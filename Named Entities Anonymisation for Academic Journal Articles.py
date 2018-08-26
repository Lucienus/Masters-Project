# import CloudConvert, python-docx and stanfordNLP Library
import cloudconvert
import docx
from collections import Counter 
import os
import re
from nltk import StanfordNERTagger
from tkinter import *
import tkinter.filedialog
from tkinter import messagebox

# change the parametre 'model_filename' to the path of 'english.conll.4class.distsim.crf.ser.gz' on your local PC
# change the parametre 'path_to_jar' to the path of 'stanford-ner.jar' on your local PC
stner = StanfordNERTagger(model_filename=r'/Users/aa/Downloads/stanford-ner-2018-02-27/classifiers/english.conll.4class.distsim.crf.ser.gz',path_to_jar=r'/Users/aa/Downloads/stanford-ner-2018-02-27/stanford-ner.jar')

# The function of conversion through CloudConvert
def convert_PDF2DOCX(path, api_key):
    api = cloudconvert.Api(api_key)
    process = api.createProcess({
        "inputformat": "pdf",
        "outputformat": "docx"
    })
    process = api.convert({
        'inputformat': 'pdf',
        'outputformat': 'docx',
        'input': 'upload',
        'file': open(path+'.pdf', 'rb')
    })
    # wait until conversion finished
    process.wait() 
    # download output file
    process.download(path+".docx") 

# The function of getting the named entities which are not 'OBJECT'
# then we can get the named entities which are 'PERSON','LOCATION' and 'ORGANIZATION'
def get_continuous_chunks(tagged_sent):
    continuous_chunk = []
    current_chunk = []

    for token, tag in tagged_sent:
        if tag != "O":
            current_chunk.append((token, tag))
        else:
            if current_chunk: # if the current chunk is not empty
                continuous_chunk.append(current_chunk)
                current_chunk = []
    # Flush the final current_chunk into the continuous_chunk, if any.
    if current_chunk:
        continuous_chunk.append(current_chunk)
    return continuous_chunk

# The function for getting the names in the previous NE list
def get_nameTags(path):
    #   read the content of Word file
    doc = docx.Document(path+".docx")

    # set a 'tagList' for storing the NEs found in the current paragraph
    tagList=[]
    
    #　set an iteration in each paragraoh
    for para in doc.paragraphs :
        for i in para.runs:
            if (i.font.superscript): # if it is a superscript, remove it 
                i.text = ''
            if (i.text.find('*') != -1): # if it is an asterisk, remove it 
                i.text = i.text.replace('*', '')
        no_Comma = para.text.replace(',',' and') #replace the commas with 'and' for reducing the errors
    #   NER processing
        tagged_sent = stner.tag(no_Comma.split())
        named_entities = get_continuous_chunks(tagged_sent)
        named_entities_str_tag = [(" ".join([token for token, tag in ne]), ne[0][1]) for ne in named_entities]

    #   eliminate the repeated tags in the list
        nameEntities = sorted(set(named_entities_str_tag),key=named_entities_str_tag.index)

    #   append the tags into 'tagList'
        i=0
        for names in nameEntities:
            if(nameEntities[i][1]=="ORGANIZATION") or (nameEntities[i][1]=="LOCATION") or (nameEntities[i][1]=="PERSON"):
                tagList.append(nameEntities[i])
                i = i+1
            else:
                i = i+1
                continue
    # before the section of 'Keywords', 'Introduction' or 'Abstract'
        if (para.text.find('KEYWORDS') != -1) or (para.text.find('Keywords') != -1) or (para.text.find('Introduction') != -1) or (para.text.find('Abstract') != -1) or (para.text.find('ABSTRACT') != -1) or (para.text.find('INTRODUCTION') != -1):
            break
        print(para.text)
    return tagList


# The function for names anonymisaiton
def anonymize_names(path):
    # read Docx file
    doc = docx.Document(path+".docx")
    # get the named entities
    tagList = get_nameTags(path)
    print(tagList)
    personList = []

    # extract all of the names
    counter=0
    for person in tagList:
        if(tagList[counter][1]=="PERSON"):
            personList.append(tagList[counter][0])
            counter = counter+1
        else:
            counter = counter+1
            continue
    print(personList)

    # step of names anonymisation
    square = ''
    for name in personList:
        for i in range(len(name)):
            square = square + '■'
        for para in doc.paragraphs:
    #       re.IGNORECAS for ignoring the uppercase or lowercase
            m = re.search(name, para.text, re.IGNORECASE)
            if bool(m):
                for i in para.runs:
                    n = re.search(name, i.text, re.IGNORECASE)
                    if bool(n):
                        text = re.sub(name, square, i.text, flags=re.IGNORECASE)
                        i.text = text
        square = ''
    # split the names into the first names and surnames
    name_split_List = []
    for element in personList:
        element = element.split()
        for e in element:
            name_split_List.append(e)

    # extract all of the names
    nameSplitList = []        
    for each in name_split_List:
        tagged_sent = stner.tag(each.split())
        named_entities = get_continuous_chunks(tagged_sent)
        named_entities_str_tag = [(" ".join([token for token, tag in ne]), ne[0][1]) for ne in named_entities]
        if named_entities_str_tag:
            if (named_entities_str_tag[0][1] == 'PERSON'):
                nameSplitList.append(each)
        else:
            continue
    print(nameSplitList)

    # step for anonymising all of the first names and surnames
    square = ''
    for name in nameSplitList:
        for i in range(len(name)):
            square = square + '■'
        for para in doc.paragraphs:
    #       re.IGNORECASE
            m = re.search(name, para.text, re.IGNORECASE)
            if bool(m):
                for i in para.runs:
                    n = re.search(name, i.text, re.IGNORECASE)
                    if bool(n):
                        text = re.sub(name, square, i.text, flags=re.IGNORECASE)
                        i.text = text
        square = ''

    doc.save(path+"-anonymous.docx")
    print('Anonymize names Finished!')

# function for anonymising the information about authors
def anonymize_info(path):
    doc = docx.Document(path+"-anonymous.docx")
    # the list for storing the superscripts
    supList = []

    # extract all superscripts before the section of 'Keywords', 'Introduction' or 'Abstract'
    for para in doc.paragraphs:
        for i in para.runs:
            if (i.font.superscript):
                i.text = i.text.replace(',','')
                supList.append(i.text)
        if (para.text.find('KEYWORDS') != -1) or (para.text.find('Keywords') != -1) or (para.text.find('Introduction') != -1) or (para.text.find('Abstract') != -1) or (para.text.find('ABSTRACT') != -1) or (para.text.find('INTRODUCTION') != -1):
                break
    supList = sorted(set(supList),key=supList.index)
    print(supList)

    counter = 0
    square = ''
    condition_flag = False  # the flag for break the loop

    # anonymise those names with the superscripts
    for para in doc.paragraphs:
        for i in para.runs:
            if (i.font.superscript) and (para.text.find('■') != -1):
                for ii in para.runs:
                    for length in range(len(ii.text)):
                        square = square + '■'
                    ii.text = square
                    square = ''
                condition_flag = True
                break
        if condition_flag:
            condition_flag = False
            break

    # anonymise the information
    for sup in supList:
        for para in doc.paragraphs:
            for i in para.runs:
                # if find a superscript and it is same as the previous superscript
                if (i.font.superscript) and (sup == i.text):
                    # NER processing
                    tagged_sent = stner.tag(para.text.split())
                    named_entities = get_continuous_chunks(tagged_sent)
                    named_entities_str_tag = [(" ".join([token for token, tag in ne]), ne[0][1]) for ne in named_entities]
                    tagList_1 = []
                    # anonymisation steps
                    square = ''
                    for tags1 in named_entities_str_tag:
                        if(tags1[1]=="ORGANIZATION") or (tags1[1]=="LOCATION") or (tags1[1]=="PERSON"):
                            tagList_1.append(tags1[0])
                        else:
                            continue
                    for element in tagList_1:
                        for length in range(len(element)):
                            square = square + '■'
                        for ii in para.runs:
                            ii.text = ii.text.replace(element,square)
                        square = ''
                    condition_flag = True
                    break
            if condition_flag:
                condition_flag = False
                break
                
    # anonymise the section of 'Acknowledgement'
    for para in doc.paragraphs:
        # if the current paragraph has the words like 'thank', 'appreaciete' or others
        thx = re.search('thank', para.text, re.IGNORECASE)
        apr = re.search('appreciate', para.text, re.IGNORECASE) or re.search('appreciation', para.text, re.IGNORECASE)
        grt = re.search('grateful', para.text, re.IGNORECASE) or re.search('correspond', para.text, re.IGNORECASE)
        if bool(thx) or bool(apr) or bool(grt):
            tagged_sent = stner.tag(para.text.split())
            named_entities = get_continuous_chunks(tagged_sent)
            named_entities_str_tag = [(" ".join([token for token, tag in ne]), ne[0][1]) for ne in named_entities]
            print(named_entities_str_tag)
            tagList_1 = []
            # anonymisation steps
            square = ''
            for tags1 in named_entities_str_tag:
                if(tags1[1]=="ORGANIZATION") or (tags1[1]=="LOCATION") or (tags1[1]=="PERSON"):
                    tagList_1.append(tags1[0])
                else:
                    continue
            print(tagList_1)
            for element in tagList_1:
                for length in range(len(element)):
                    square = square + '■'
                for i in para.runs:
                    i.text = i.text.replace(element,square)
                square = ''

    doc.save(path+"-anonymous_2.docx")
    print('Finished!')


# The class of jupming out a window
class myWindow:
    def __init__(self, root, myTitle, flag, path):
        self.top = tkinter.Toplevel(root, width=100, height=100)
        self.top.title(myTitle)
        self.top.attributes('-topmost', 1)
        if flag == 1:
            S = Scrollbar(self.top)
            t = Text(self.top ,height=100, width=100)
            S.pack(side = RIGHT, fill = Y)
            t.pack(side = LEFT, fill = Y)
            S.config(command = t.yview)
            t.config(yscrollcommand = S.set)
            doc = docx.Document(path+"-anonymous_2.docx")
            text = ''
            for para in doc.paragraphs:
                text = text + para.text + '\n'
            t.insert(END, text)
            
# the function for closing the window
def on_closing():
    if messagebox.askokcancel("Quit", "Do you want to quit?"):
        root.destroy()

# the function for choosing the input PDF file
def choose():
    filename = tkinter.filedialog.askopenfilename()
    if filename != '':
        lb.config(text = "The file you selected is："+filename)
        global file_path
        file_path = filename
        file_path = file_path.replace('.pdf','')
        print(file_path)
    else:
        lb.config(text = "You didn't select any file.")
        
# the function for starting processing
def run():
    api = api_t.get()
    convert_PDF2DOCX(file_path, api)
    anonymize_names(file_path)
    anonymize_info(file_path)
    messagebox.showinfo(title='Done',message='Program Finished!')
    if window1.get()==0:
        window1.set(1)
        w1 = myWindow(root, 'Output Window', 1, file_path)
        window1.set(0)
        
# the function for converting docx to pdf
def convert_d2p():
    api = api_t.get()
    convert_DOCX2PDF(file_path, api)
    messagebox.showinfo(title='Done',message='Converting successfully!')
        
# the function of the button for other content anonymisation
def anonymise_other():
    doc = docx.Document(file_path + "-anonymous_2.docx")
    text = t1.get()
    square = ''
    for para in doc.paragraphs:
        find = re.search(text, para.text, re.IGNORECASE)
        if bool(find):
            for i in para.runs:
                for length in range(len(text)):
                    square = square + '■'
                i.text = i.text.replace(text,square)
                square = ''
    doc.save(file_path + "-anonymous_2.docx")
    messagebox.showinfo(title='Done',message='Anonymised!')
    if window1.get()==0:
        window1.set(1)
        w1 = myWindow(root, 'Output Window', 1, file_path)
        window1.set(0)

# the function for converting pdf to docx
def convert_DOCX2PDF(path, api_key):
    api = cloudconvert.Api(api_key)
    process = api.createProcess({
        "inputformat": "docx",
        "outputformat": "pdf"
    })
    process = api.convert({
        'inputformat': 'docx',
        'outputformat': 'pdf',
        'input': 'upload',
        'file': open(path + '-anonymous_2.docx', 'rb')
    })
    process.wait() # wait until conversion finished
    process.download(path + "-anonymised.pdf") # download output file

# The main function for building the GUI
if __name__=="__main__":
    
    root = Tk()
    root.protocol("WM_DELETE_WINDOW", on_closing)
    root.geometry('800x300+400+200')

    lb1 = Label(root,text = 'Your API:')
    lb1.pack()
    entryTxt = StringVar()
    entryTxt.set('ry91t9nQMBqGnZBFCzuk09pxcKbvt1ulG9joUwHD4Et8xNGAZLtjSyifYlqDgH9q')
    global api_t
    api_t = Entry(root,bg='#cacaca',textvariable = entryTxt,width=57)
    api_t.pack()


    lb = Label(root,text = '')
    lb.pack()
    btn1 = Button(root,text="Choose file",command=choose, width = 15)
    btn1.pack()
    btn2 = Button(root,text="RUN",command=run, width = 15)
    btn2.pack()
    
    lb2 = Label(root,text = 'ARE YOU SATISFIED WITH THE OUTPUT?\nIF NOT, PLEASE PASTE THE CONTENT YOU STILL WANT TO ANONYMISE:')
    lb2.pack()
    global t1
    t1 = Entry(root,bg='#cacaca',width=57)
    t1.pack()
    btn3 = Button(root,text="ANONYMISE",command=anonymise_other, width = 15)
    btn3.pack()
    
    btn4 = Button(root,text="GENERATE PDF",command=convert_d2p, width = 15)
    btn4.pack()
    
    global window1
    window1 = tkinter.IntVar(root, value=0)
    root.mainloop()